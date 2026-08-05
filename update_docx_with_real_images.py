import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import re

# Map figure key phrases to image filenames in extracted_images and project root
FIGURE_IMAGE_MAP = {
    'Fig 3.1.1': ['extracted_images/image6.png'],
    'Fig 3.1.2': ['extracted_images/image25.png'],
    'Fig 4.2.1': ['extracted_images/image18.png'],
    'Fig 4.3.1': ['extracted_images/image2.png'],
    'Fig 4.3.2': ['extracted_images/image2.png'],
    'Fig 4.3.3': ['extracted_images/image8.png', 'extracted_images/image10.png'],
    'Fig 4.3.4': ['extracted_images/image27.png'],
    'Fig 4.3.5': ['extracted_images/image22.png'],
    'Fig 4.3.7': ['extracted_images/image14.png'],
    'Fig 4.3.8': ['extracted_images/image24.png'],
    'Fig 4.3.9': ['extracted_images/image1.png', 'extracted_images/image3.png', 'extracted_images/image19.png'],
    'Fig 4.3.10': ['extracted_images/image15.jpg'],
    'Fig 2.1': ['extracted_images/image26.png'],
    'Fig 2.2': ['extracted_images/image12.png'],
    'Fig 2.3': ['extracted_images/image5.png', 'extracted_images/image13.png'],
    'Fig 6.1': ['extracted_images/image20.png'],
    'Fig 9.1': ['extracted_images/image17.gif'],
    'Fig 9.2': ['extracted_images/image16.gif'],
    'Fig 9.3': ['extracted_images/image21.png'],
    'Fig 9.4': ['extracted_images/image11.jpg'],
    'Fig 9.5': ['extracted_images/image9.jpg'],
    'Fig 9.6': ['extracted_images/image7.jpg'],
    'Fig 9.7': ['HomeView.png', 'FrontView.png'],
    'Fig 9.8': ['RightView.png', 'LeftView.png', 'BottomView.png', 'BackView.png'],
}

def create_rollopod_docx_with_images(md_filepath, docx_filepath):
    doc = Document()
    
    # Page Setup - Margins 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    with open(md_filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_lines = []
    
    def process_inline(paragraph, text):
        pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)')
        tokens = pattern.split(text)
        for token in tokens:
            if not token:
                continue
            if token.startswith('**') and token.endswith('**'):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith('*') and token.endswith('*'):
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            elif token.startswith('`') and token.endswith('`'):
                run = paragraph.add_run(token[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            else:
                paragraph.add_run(token)

    def render_table(table_data):
        if not table_data:
            return
        rows = [row for row in table_data if not all(c == '-' or c == ':' or c == '---' for c in row)]
        if not rows:
            return
        
        num_rows = len(rows)
        num_cols = max(len(r) for r in rows)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for idx, text in enumerate(rows[0]):
            table.rows[0].cells[idx].text = text.strip()
            shading_elm = parse_xml(r'<w:shd {} w:fill="1F497D"/>'.format(nsdecls('w')))
            table.rows[0].cells[idx]._tc.get_or_add_tcPr().append(shading_elm)
            for p in table.rows[0].cells[idx].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.font.name = 'Calibri'
                    
        for r_idx in range(1, num_rows):
            row_cells = table.rows[r_idx].cells
            row_data = rows[r_idx]
            for c_idx in range(min(len(row_data), num_cols)):
                cell_text = row_data[c_idx].strip()
                row_cells[c_idx].text = cell_text
                if r_idx % 2 == 0:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="F2F5F8"/>'.format(nsdecls('w')))
                    row_cells[c_idx]._tc.get_or_add_tcPr().append(shading_elm)
                for p in row_cells[c_idx].paragraphs:
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(3)
                    for r in p.runs:
                        r.font.name = 'Calibri'
                        r.font.size = Pt(10)
                        if r_idx == num_rows - 1:
                            r.font.bold = True
                            
        doc.add_paragraph()

    in_code_block = False
    
    for line in lines:
        raw_line = line
        line_str = line.strip()
        
        if line_str.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(raw_line.rstrip('\n'))
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        if '|' in line_str and not line_str.startswith('#'):
            parts = [p.strip() for p in line_str.split('|')]
            if parts and parts[0] == '':
                parts = parts[1:]
            if parts and parts[-1] == '':
                parts = parts[:-1]
            if len(parts) > 1:
                if all(re.match(r'^[:\- ]+$', cell) for cell in parts):
                    continue
                if not in_table:
                    in_table = True
                    table_lines = []
                table_lines.append(parts)
                continue
        else:
            if in_table:
                render_table(table_lines)
                in_table = False
                table_lines = []

        if not line_str:
            continue
            
        # Headers
        if line_str.startswith('# '):
            title_text = line_str[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(title_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif line_str.startswith('## '):
            h_text = line_str[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(h_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif line_str.startswith('### '):
            h_text = line_str[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(h_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(12.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif line_str.startswith('---'):
            continue
        elif line_str.startswith('- ') or line_str.startswith('* '):
            bullet_text = line_str[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            process_inline(p, bullet_text)
        elif re.match(r'^\d+\.\s', line_str):
            item_text = re.sub(r'^\d+\.\s', '', line_str).strip()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            process_inline(p, item_text)
        elif line_str.startswith('Fig ') or line_str.startswith('Figure '):
            # Embed Image(s) if mapped
            matched_key = None
            for key in FIGURE_IMAGE_MAP:
                if key in line_str:
                    matched_key = key
                    break
            
            if matched_key and matched_key in FIGURE_IMAGE_MAP:
                img_paths = FIGURE_IMAGE_MAP[matched_key]
                for img_p in img_paths:
                    if os.path.exists(img_p):
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_para.paragraph_format.space_before = Pt(8)
                        img_para.paragraph_format.space_after = Pt(4)
                        try:
                            img_para.add_run().add_picture(img_p, width=Inches(4.8))
                        except Exception as e:
                            print(f"Error embedding image {img_p}: {e}")
            
            # Caption Paragraph
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(8)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line_str)
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            process_inline(p, line_str)

    if in_table:
        render_table(table_lines)
        
    doc.save(docx_filepath)
    print(f"Successfully generated {docx_filepath} with all embedded render images.")

if __name__ == '__main__':
    create_rollopod_docx_with_images('HexapodTheoriticalIdeation.md', 'HexapodTheoriticalIdeation.docx')
