import subprocess
import docx
import zipfile
import os

# Extract original docx from git HEAD
with open('original.docx', 'wb') as f:
    subprocess.run(['git', 'show', 'HEAD:HexapodTheoriticalIdeation.docx'], stdout=f)

print("Original docx extracted.")

# Extract all images from original.docx
os.makedirs('extracted_images', exist_ok=True)
with zipfile.ZipFile('original.docx', 'r') as z:
    for member in z.namelist():
        if member.startswith('word/media/'):
            filename = os.path.basename(member)
            target_path = os.path.join('extracted_images', filename)
            with open(target_path, 'wb') as out_f:
                out_f.write(z.read(member))
            print(f"Extracted image: {filename}")

doc = docx.Document('original.docx')
print(f"Total paragraphs in original docx: {len(doc.paragraphs)}")

# List paragraphs with image drawings and captions
for i, p in enumerate(doc.paragraphs):
    has_drawing = 'drawing' in p._element.xml
    if has_drawing or 'Fig' in p.text or 'fig' in p.text or 'FIG' in p.text:
        print(f"P{i:03d} | Has Drawing: {str(has_drawing):5s} | Text: {p.text.strip()}")
